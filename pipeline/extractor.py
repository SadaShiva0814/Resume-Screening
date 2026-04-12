"""
Text Extraction Engine
Extracts clean text from multiple resume formats:
  - PDF (native text-selectable)
  - PDF (image-based / scanned) via OCR
  - DOCX
  - TXT
  
Uses a smart fallback strategy: tries native extraction first,
falls back to OCR if the PDF appears to be image-based.
"""

import os
import re
import logging
from pathlib import Path

logger = logging.getLogger(__name__)


def extract_text(file_path):
    """
    Extract text from any supported resume format.
    
    Smart detection:
    1. Check file extension
    2. For PDFs: try native extraction first
    3. If native extraction yields < 50 chars, fall back to OCR
    
    Returns: cleaned text string
    """
    file_path = str(file_path)
    ext = Path(file_path).suffix.lower()
    
    if ext == '.txt':
        return _extract_txt(file_path)
    elif ext == '.docx':
        return _extract_docx(file_path)
    elif ext in ('.pdf',):
        return _extract_pdf_smart(file_path)
    else:
        logger.warning(f"Unsupported file format: {ext}")
        return ""


def _extract_txt(file_path):
    """Extract text from plain text file."""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()
        return _clean_text(text)
    except Exception as e:
        logger.error(f"TXT extraction failed for {file_path}: {e}")
        return ""


def _extract_docx(file_path):
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document
        doc = Document(file_path)
        
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        
        text = '\n'.join(paragraphs)
        return _clean_text(text)
    except Exception as e:
        logger.error(f"DOCX extraction failed for {file_path}: {e}")
        return ""


def _extract_pdf_smart(file_path):
    """
    Smart PDF extraction:
    1. Try PyMuPDF (fitz) for native text — fastest and most accurate
    2. If result is too short, try PyPDF2 as backup
    3. If still too short, assume image-based PDF and use OCR
    """
    # Strategy 1: PyMuPDF (best native PDF parser)
    text = _extract_pdf_pymupdf(file_path)
    
    if len(text.strip()) >= 50:
        return _clean_text(text)
    
    # Strategy 2: PyPDF2 fallback
    text = _extract_pdf_pypdf2(file_path)
    
    if len(text.strip()) >= 50:
        return _clean_text(text)
    
    # Strategy 3: OCR (image-based PDF)
    logger.info(f"Native extraction yielded <50 chars, using OCR for: {file_path}")
    text = _extract_pdf_ocr(file_path)
    return _clean_text(text)


def _extract_pdf_pymupdf(file_path):
    """Extract text using PyMuPDF (fitz) — fastest native parser."""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(file_path)
        text_parts = []
        for page in doc:
            text_parts.append(page.get_text())
        doc.close()
        return '\n'.join(text_parts)
    except Exception as e:
        logger.error(f"PyMuPDF extraction failed for {file_path}: {e}")
        return ""


def _extract_pdf_pypdf2(file_path):
    """Extract text using PyPDF2 — backup native parser."""
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        return '\n'.join(text_parts)
    except Exception as e:
        logger.error(f"PyPDF2 extraction failed for {file_path}: {e}")
        return ""


def _extract_pdf_ocr(file_path):
    """Extract text from image-based PDF using Tesseract OCR."""
    try:
        from pdf2image import convert_from_path
        import pytesseract
        from config import Config
        
        # Set tesseract command path
        pytesseract.pytesseract.tesseract_cmd = Config.TESSERACT_CMD
        
        # Convert PDF pages to images
        images = convert_from_path(file_path, dpi=300)
        
        text_parts = []
        for i, image in enumerate(images):
            # OCR each page
            page_text = pytesseract.image_to_string(image, lang='eng')
            if page_text.strip():
                text_parts.append(page_text)
        
        return '\n'.join(text_parts)
    except Exception as e:
        logger.error(f"OCR extraction failed for {file_path}: {e}")
        return ""


def _clean_text(text):
    """
    Clean extracted text:
    - Remove excessive whitespace
    - Remove special characters that break NLP
    - Normalize line breaks
    - Remove page numbers and headers/footers
    """
    if not text:
        return ""
    
    # Remove null bytes and control characters
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)
    
    # Normalize unicode
    text = text.encode('ascii', 'ignore').decode('ascii')
    
    # Remove excessive whitespace but preserve paragraph structure
    text = re.sub(r'[ \t]+', ' ', text)           # Multiple spaces/tabs to single space
    text = re.sub(r'\n{3,}', '\n\n', text)         # Max 2 consecutive newlines
    text = re.sub(r'^\s+$', '', text, flags=re.MULTILINE)  # Remove blank lines with only spaces
    
    # Remove common PDF artifacts
    text = re.sub(r'Page \d+ of \d+', '', text, flags=re.IGNORECASE)
    text = re.sub(r'\f', '\n', text)               # Form feeds to newlines
    
    return text.strip()


def extract_text_from_string(raw_text):
    """Clean pre-extracted text (e.g., from Kaggle CSV)."""
    return _clean_text(raw_text)


def extract_text_from_html(html_text):
    """
    Extract structured text from HTML resume (Kaggle Resume_html column).
    
    The HTML has semantic div classes like:
    - SECTION_SUMM (Summary)
    - SECTION_EXPR (Experience)  
    - SECTION_EDUC (Education)
    - SECTION_SKLL (Skills)
    - SECTION_NAME (Name/Header)
    - SECTION_HILT (Highlights)
    - SECTION_PROJ (Projects)
    - SECTION_CERT (Certifications)
    
    We parse these to extract STRUCTURED section text.
    """
    if not html_text or not isinstance(html_text, str):
        return ""
    
    try:
        # Simple HTML tag stripping with structure preservation
        import re
        
        # Insert newlines before section divs for better splitting
        text = re.sub(r'<div[^>]*class="[^"]*section[^"]*"[^>]*>', '\n\n', html_text, flags=re.IGNORECASE)
        text = re.sub(r'<div[^>]*class="[^"]*heading[^"]*"[^>]*>', '\n', text, flags=re.IGNORECASE)
        text = re.sub(r'<div[^>]*class="[^"]*sectiontitle[^"]*"[^>]*>', '\n', text, flags=re.IGNORECASE)
        text = re.sub(r'<div[^>]*class="[^"]*paragraph[^"]*"[^>]*>', '\n', text, flags=re.IGNORECASE)
        
        # Replace br and p tags with newlines
        text = re.sub(r'<br\s*/?>', '\n', text, flags=re.IGNORECASE)
        text = re.sub(r'</p>', '\n', text, flags=re.IGNORECASE)
        text = re.sub(r'</div>', '\n', text, flags=re.IGNORECASE)
        text = re.sub(r'</li>', '\n', text, flags=re.IGNORECASE)
        
        # Strip all remaining HTML tags
        text = re.sub(r'<[^>]+>', ' ', text)
        
        # Decode HTML entities
        text = text.replace('&amp;', '&')
        text = text.replace('&lt;', '<')
        text = text.replace('&gt;', '>')
        text = text.replace('&nbsp;', ' ')
        text = text.replace('&#8226;', '•')
        text = text.replace('&#39;', "'")
        text = text.replace('&quot;', '"')
        
        return _clean_text(text)
    except Exception as e:
        logger.error(f"HTML extraction failed: {e}")
        return _clean_text(re.sub(r'<[^>]+>', ' ', html_text))


def extract_sections_from_html(html_text):
    """
    Extract structured sections directly from Kaggle HTML.
    Returns a dict mapping section names to their text content.
    
    This is MUCH better than parsing the flat Resume_str because
    the HTML preserves the original section structure.
    """
    if not html_text or not isinstance(html_text, str):
        return {}
    
    import re
    
    # Map HTML section IDs to our section names
    section_map = {
        'SUMM': 'summary',
        'EXPR': 'experience',
        'EDUC': 'education',
        'SKLL': 'skills',
        'HILT': 'skills',      # Highlights → Skills
        'NAME': 'header',
        'PROJ': 'projects',
        'CERT': 'education',    # Certifications → Education
        'ACCM': 'projects',     # Accomplishments → Projects
        'MISC': 'summary',
        'ADDL': 'summary',     # Additional → Summary
    }
    
    sections = {}
    
    # Find all section divs — section code is 4 letters followed by digits
    section_pattern = re.compile(
        r'<div[^>]*id="SECTION_([A-Z]{4})\d+"[^>]*>(.*?)(?=<div[^>]*id="SECTION_[A-Z]{4}\d+|$)',
        re.DOTALL | re.IGNORECASE
    )
    
    for match in section_pattern.finditer(html_text):
        section_code = match.group(1)
        section_html = match.group(2)
        
        # Map to our section name
        section_name = section_map.get(section_code, 'summary')
        
        if section_name == 'header':
            # Extract the header/title text (job title from Kaggle, or name)
            header_text = re.sub(r'<[^>]+>', ' ', section_html)
            header_text = re.sub(r'\s+', ' ', header_text).strip()
            header_text = header_text.replace('&amp;', '&').replace('&nbsp;', ' ')
            if header_text:
                sections['header'] = header_text
            continue  # Don't add header to content sections
        
        # Extract clean text from the section HTML
        # Get section title if present
        title_match = re.search(r'class="sectiontitle"[^>]*>([^<]+)', section_html)
        
        # Strip HTML tags
        text = re.sub(r'<[^>]+>', ' ', section_html)
        text = re.sub(r'\s+', ' ', text).strip()
        
        # Decode entities 
        text = text.replace('&amp;', '&').replace('&nbsp;', ' ')
        text = text.replace('&#8226;', '•').replace('&#39;', "'")
        
        if text and len(text) > 10:
            if section_name in sections:
                sections[section_name] += '\n' + text
            else:
                sections[section_name] = text
    
    return sections
